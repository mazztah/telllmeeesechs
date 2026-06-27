# dv.py – Universal Data Processor (FINAL + kompatibel mit main.py)
import logging
import os
import mimetypes
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
from pydub import AudioSegment
from docx import Document
import PyPDF2
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import json

logger = logging.getLogger(__name__)

def get_mime(file_path: str) -> str:
    mime, _ = mimetypes.guess_type(file_path)
    if mime:
        return mime
    ext = os.path.splitext(file_path)[1].lower()
    fallback = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.mp3': 'audio/mpeg',
        '.wav': 'audio/wav',
        '.mp4': 'video/mp4',
        '.txt': 'text/plain',
        '.csv': 'text/csv',
        '.json': 'application/json',
    }
    return fallback.get(ext, "application/octet-stream")


# ====================== LESEN ======================
def extract_content(file_path: str, max_chars: int = 12000) -> str:
    mime = get_mime(file_path)
    name = os.path.basename(file_path)

    try:
        if mime.startswith("image/"):
            img = Image.open(file_path)
            return f"📸 Bild '{name}': {img.format} | Größe: {img.size} | Modus: {img.mode}"

        elif mime.startswith("audio/"):
            audio = AudioSegment.from_file(file_path)
            return f"🎵 Audio '{name}': {len(audio)/1000:.1f}s | Kanäle: {audio.channels} | {audio.frame_rate}Hz"

        elif mime.startswith("video/"):
            return f"🎥 Video '{name}': Datei erkannt"

        elif mime == "application/pdf":
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = "".join(p.extract_text() or "" for p in reader.pages)
            return f"📄 PDF '{name}' ({len(reader.pages)} Seiten):\n{text[:max_chars]} [...]"

        elif "wordprocessingml" in mime or mime.endswith("document"):
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return f"📝 Word '{name}':\n{text[:max_chars]} [...]"

        elif mime in ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "text/csv"):
            df = pd.read_excel(file_path) if "excel" in mime else pd.read_csv(file_path)
            return f"📊 {mime.split('/')[-1].upper()} '{name}' ({df.shape[0]} Zeilen):\n{df.head(10).to_string()}"

        elif mime == "application/json":
            with open(file_path, encoding="utf-8") as f:
                data = json.load(f)
            return f"📋 JSON '{name}':\n{json.dumps(data, indent=2, ensure_ascii=False)[:max_chars]}"

        elif mime.startswith("text/"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read(max_chars)
            return f"📜 Text '{name}':\n{text}"

        return f"❓ Unbekanntes Format '{name}' (MIME: {mime})"

    except Exception as e:
        logger.error(f"extract_content Fehler bei {name}: {e}")
        return f"💥 Extract-Fehler bei {name}: {str(e)[:200]}"


# ====================== ALTE FUNKTIONEN (für main.py Kompatibilität) ======================
def resize_image(file_path: str, size: tuple = (800, 800)) -> BytesIO:
    img = Image.open(file_path).resize(size, Image.Resampling.LANCZOS)
    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def trim_audio(file_path: str, start: float = 0.0, end: float | None = None) -> BytesIO:
    audio = AudioSegment.from_file(file_path)
    if end is None:
        end = len(audio) / 1000
    trimmed = audio[int(start * 1000):int(end * 1000)]
    buffer = BytesIO()
    trimmed.export(buffer, format="mp3")
    buffer.seek(0)
    return buffer


def save_locally(buffer: BytesIO, filename: str, folder: str = "/tmp/processed") -> str:
    os.makedirs(folder, exist_ok=True)
    path = os.path.join(folder, filename)
    with open(path, "wb") as f:
        f.write(buffer.getvalue())
    return path


# ====================== NEUE ERSTELL-FUNKTIONEN ======================
def create_pdf_from_text(text: str, title: str = "Generated.pdf") -> BytesIO:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    y = 750
    for line in text.split("\n"):
        if y < 50:
            c.showPage()
            y = 750
        c.drawString(50, y, line[:90])
        y -= 15
    c.save()
    buffer.seek(0)
    return buffer


def create_docx_from_text(text: str, title: str = "Generated.docx") -> BytesIO:
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_excel_from_data(data: list[list], columns: list[str], title: str = "Generated.xlsx") -> BytesIO:
    """Erstellt ein formatiertes XLSX mit Hyperlinks für URL-Spalten."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = "Ergebnisse"
        ws.freeze_panes = "A2"  # Header einfrieren

        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill("solid", fgColor="6D28D9")
        header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        url_font = Font(color="0563C1", underline="single", size=10)
        cell_align = Alignment(vertical="top", wrap_text=True)
        thin = Side(style="thin", color="E5E7EB")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        # Spaltenbreiten je nach Typ
        col_widths = {
            "URL": 55, "Beschreibung (Snippet)": 55, "Titel": 38,
            "Firma": 28, "Ort": 22, "Job-ID": 14,
        }

        # Header-Zeile
        for ci, col_name in enumerate(columns, 1):
            cell = ws.cell(row=1, column=ci, value=col_name)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = border
            letter = get_column_letter(ci)
            ws.column_dimensions[letter].width = col_widths.get(col_name, 20)

        ws.row_dimensions[1].height = 22

        # Datenzeilen
        for ri, row in enumerate(data, 2):
            for ci, val in enumerate(row, 1):
                col_name = columns[ci - 1] if ci - 1 < len(columns) else ""
                str_val = str(val) if val is not None else ""
                cell = ws.cell(row=ri, column=ci, value=str_val)
                cell.alignment = cell_align
                cell.border = border
                # URL → Hyperlink
                if col_name == "URL" and str_val.startswith("http"):
                    cell.hyperlink = str_val
                    cell.value = str_val
                    cell.font = url_font
                else:
                    cell.font = Font(size=10)
                # Zebra-Streifen
                if ri % 2 == 0:
                    cell.fill = PatternFill("solid", fgColor="F5F3FF")

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer

    except ImportError:
        # Fallback auf pandas falls openpyxl fehlt
        df = pd.DataFrame(data, columns=columns)
        buffer = BytesIO()
        df.to_excel(buffer, index=False, engine="openpyxl")
        buffer.seek(0)
        return buffer



def create_jobqueen_excel(jobs: list, queries: list = None, export_date: str = "") -> "BytesIO":
    """
    Erstellt professionelles JobQueen-Export-XLSX.
    Sheet 1: Zusammenfassung | Sheet 2: Ausgewählte Stellen
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import datetime as _dt
    from collections import Counter

    wb   = Workbook()
    thin = Side(style="thin",   color="C4B5FD")
    med  = Side(style="medium", color="7C3AED")
    brd  = Border(left=thin, right=thin, top=thin, bottom=thin)
    brdH = Border(left=med,  right=med,  top=med,  bottom=med)

    PURPLE="#6D28D9"; PURLT="#EDE9FE"; GOLD="#D4AF37"
    WHITE="#FFFFFF";  DARK="#1E1B4B"; DARK2="#312E81"
    GRAY="#6B7280";   GREEN="#065F46"; GBGG="#D1FAE5"; LINK="0563C1"

    def H(ws,r,c,v,bg=PURPLE,fg=WHITE,sz=11,bold=True,ctr=True):
        x=ws.cell(row=r,column=c,value=v)
        x.font=Font(bold=bold,color=fg,size=sz,name="Calibri")
        x.fill=PatternFill("solid",fgColor=bg)
        x.alignment=Alignment(horizontal="center" if ctr else "left",vertical="center",wrap_text=True)
        x.border=brdH; return x

    def D(ws,r,c,v,bold=False,color=DARK,bg=None,url=None,ctr=False,sz=10):
        s=str(v) if v is not None else ""
        x=ws.cell(row=r,column=c,value=s)
        x.font=Font(bold=bold,color=color,size=sz,name="Calibri",underline="single" if url else None)
        x.alignment=Alignment(horizontal="center" if ctr else "left",vertical="top",wrap_text=True)
        x.border=brd
        if bg: x.fill=PatternFill("solid",fgColor=bg)
        if url and str(url).startswith("http"): x.hyperlink=str(url); x.value=str(v) if v else str(url)
        return x

    # ── Sheet 1: Zusammenfassung ──────────────────────────────────────────
    ws1=wb.active; ws1.title="Zusammenfassung"; ws1.sheet_view.showGridLines=False
    ws1.merge_cells("A1:F1")
    c=ws1["A1"]; c.value="👑  JobQueen  –  Stellensuche Export"
    c.font=Font(bold=True,color=GOLD,size=18,name="Calibri")
    c.fill=PatternFill("solid",fgColor=DARK); c.alignment=Alignment(horizontal="center",vertical="center")
    ws1.row_dimensions[1].height=38

    ws1.merge_cells("A2:F2")
    c=ws1["A2"]; ed=export_date or _dt.datetime.now().strftime("%d.%m.%Y %H:%M")
    c.value=f"Exportiert am {ed}   ·   {len(jobs)} Stelle(n) ausgewählt"
    c.font=Font(color=WHITE,size=10,italic=True,name="Calibri")
    c.fill=PatternFill("solid",fgColor=DARK2); c.alignment=Alignment(horizontal="center",vertical="center")
    ws1.row_dimensions[2].height=20; ws1.row_dimensions[3].height=8

    H(ws1,4,1,"Suchanfrage",bg=PURPLE,sz=10); H(ws1,4,2,"Plattform",bg=PURPLE,sz=10); H(ws1,4,3,"Stellen",bg=PURPLE,sz=10)
    src_counts=Counter(j.get("source") or "–" for j in jobs)
    qlist=list(dict.fromkeys([q for q in (queries or []) if q]))
    n=max(len(qlist),len(src_counts),1)
    for i in range(n):
        r=5+i; bg=PURLT if i%2==0 else WHITE
        D(ws1,r,1,qlist[i] if i<len(qlist) else "",bg=bg)
    for i,(src,cnt) in enumerate(src_counts.items()):
        r=5+i; bg=PURLT if i%2==0 else WHITE
        D(ws1,r,2,src,bg=bg); D(ws1,r,3,str(cnt),bg=bg,ctr=True)
    tr=5+n+1; ws1.merge_cells(f"A{tr}:B{tr}")
    H(ws1,tr,1,f"✅  Gesamt: {len(jobs)} Stellen ausgewählt",bg=GREEN,sz=11); D(ws1,tr,3,"",bg=GBGG)
    for col,w in zip("ABCDEF",[44,28,14,10,10,10]):
        ws1.column_dimensions[col].width=w

    # ── Sheet 2: Ausgewählte Stellen ──────────────────────────────────────
    ws2=wb.create_sheet("Ausgewählte Stellen"); ws2.sheet_view.showGridLines=False; ws2.freeze_panes="A3"
    COLS=[("Nr.",6),("Jobtitel",38),("Firma",28),("Ort",20),("Quelle",20),("Datum",16),("Link zur Stelle",52),("Kurzbeschreibung",54)]
    n2=len(COLS)
    ws2.merge_cells(f"A1:{get_column_letter(n2)}1")
    c=ws2["A1"]; c.value=f"👑  JobQueen  –  {len(jobs)} ausgewählte Stellen"
    c.font=Font(bold=True,color=GOLD,size=14,name="Calibri")
    c.fill=PatternFill("solid",fgColor=DARK); c.alignment=Alignment(horizontal="center",vertical="center")
    ws2.row_dimensions[1].height=30
    for ci,(name,w) in enumerate(COLS,1):
        H(ws2,2,ci,name,bg=PURPLE,sz=10); ws2.column_dimensions[get_column_letter(ci)].width=w
    ws2.row_dimensions[2].height=22

    for ri,j in enumerate(jobs,3):
        bg=PURLT if (ri-3)%2==0 else WHITE
        url=j.get("url") or ""
        raw_d=j.get("date") or ""
        try: date_str=_dt.datetime.fromisoformat(raw_d.split("T")[0]).strftime("%d.%m.%Y")
        except: date_str=raw_d
        D(ws2,ri,1,ri-2,bg=bg,bold=True,color=PURPLE,ctr=True,sz=10)
        D(ws2,ri,2,j.get("title") or "",bg=bg,bold=True,sz=10)
        D(ws2,ri,3,j.get("company") or "",bg=bg,sz=10)
        D(ws2,ri,4,j.get("location") or "",bg=bg,sz=10)
        D(ws2,ri,5,j.get("source") or "",bg=bg,color=GRAY,sz=9)
        D(ws2,ri,6,date_str,bg=bg,ctr=True,sz=10)
        D(ws2,ri,7,"🔗 Zur Stelle öffnen",bg=bg,url=url,color=LINK,bold=True,sz=10)
        D(ws2,ri,8,(j.get("description_snippet") or "")[:300],bg=bg,color=GRAY,sz=9)
        ws2.row_dimensions[ri].height=44

    buf=BytesIO(); wb.save(buf); buf.seek(0); return buf


def create_chart_from_df(df: pd.DataFrame, title: str = "Chart") -> BytesIO:
    img = Image.new("RGB", (900, 600), "#2a0044")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 24)
    except:
        font = ImageFont.load_default()

    draw.text((30, 30), title, fill="#ffccff", font=font)

    if len(df.columns) > 1:
        values = df.iloc[:, 1].head(8)
        max_val = max(values) if len(values) > 0 else 1
        for i, val in enumerate(values):
            x = 80 + i * 90
            height = int(380 * (val / max_val))
            draw.rectangle([x, 500 - height, x + 65, 500], fill="#ff66cc")
            draw.text((x + 10, 510), str(val), fill="white", font=font)

    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def create_cv_excel(profile: dict) -> BytesIO:
    """Exportiert CV-Analyse-Ergebnisse als formatiertes XLSX."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    thin = Side(style="thin", color="D1D5DB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def hdr_cell(ws, row, col, value, fg="6D28D9"):
        c = ws.cell(row=row, column=col, value=value)
        c.font = Font(bold=True, color="FFFFFF", size=10)
        c.fill = PatternFill("solid", fgColor=fg)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = border
        return c

    def data_cell(ws, row, col, value, bold=False, color=None):
        c = ws.cell(row=row, column=col, value=str(value) if value is not None else "")
        c.font = Font(bold=bold, color=color or "1F2937", size=10)
        c.alignment = Alignment(vertical="top", wrap_text=True)
        c.border = border
        if row % 2 == 0:
            c.fill = PatternFill("solid", fgColor="F5F3FF")
        return c

    # ── Sheet 1: Übersicht ──────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Übersicht"
    ws1.freeze_panes = "A2"
    ws1.column_dimensions["A"].width = 28
    ws1.column_dimensions["B"].width = 55

    hdr_cell(ws1, 1, 1, "Feld")
    hdr_cell(ws1, 1, 2, "Wert")

    rows_overview = [
        ("Name",            profile.get("name") or "—"),
        ("Berufserfahrung", f"{profile.get('experience_years', 0)} Jahre {profile.get('experience_months', 0)} Monate"),
        ("Skills",         " · ".join(profile.get("skills") or []) or "—"),
        ("Sprachen",       " · ".join(profile.get("languages") or []) or "—"),
    ]
    for ri, (k, v) in enumerate(rows_overview, 2):
        data_cell(ws1, ri, 1, k, bold=True)
        data_cell(ws1, ri, 2, v)

    # ── Sheet 2: Stärken ───────────────────────────────────────────
    ws2 = wb.create_sheet("Stärken")
    ws2.freeze_panes = "A2"
    ws2.column_dimensions["A"].width = 30
    ws2.column_dimensions["B"].width = 50
    ws2.column_dimensions["C"].width = 45

    hdr_cell(ws2, 1, 1, "Stärke")
    hdr_cell(ws2, 1, 2, "Beleg aus dem Lebenslauf")
    hdr_cell(ws2, 1, 3, "Relevanz")

    for ri, s in enumerate(profile.get("strengths") or [], 2):
        data_cell(ws2, ri, 1, s.get("strength") or "", bold=True)
        data_cell(ws2, ri, 2, s.get("evidence") or "")
        data_cell(ws2, ri, 3, s.get("relevance") or "")

    # ── Sheet 3: Empfohlene Jobtitel ───────────────────────────────
    ws3 = wb.create_sheet("Empfohlene Jobtitel")
    ws3.freeze_panes = "A2"
    ws3.column_dimensions["A"].width = 38
    ws3.column_dimensions["B"].width = 60

    hdr_cell(ws3, 1, 1, "Jobtitel")
    hdr_cell(ws3, 1, 2, "Begründung")

    for ri, jt in enumerate(profile.get("suggested_job_titles") or [], 2):
        data_cell(ws3, ri, 1, jt.get("title") or "", bold=True)
        data_cell(ws3, ri, 2, jt.get("reason") or "")

    # ── Sheet 4: Berufserfahrung ────────────────────────────────────
    ws4 = wb.create_sheet("Berufserfahrung")
    ws4.freeze_panes = "A2"
    ws4.column_dimensions["A"].width = 32
    ws4.column_dimensions["B"].width = 30
    ws4.column_dimensions["C"].width = 14
    ws4.column_dimensions["D"].width = 14
    ws4.column_dimensions["E"].width = 12

    hdr_cell(ws4, 1, 1, "Position")
    hdr_cell(ws4, 1, 2, "Unternehmen")
    hdr_cell(ws4, 1, 3, "Von")
    hdr_cell(ws4, 1, 4, "Bis")
    hdr_cell(ws4, 1, 5, "Monate")

    roles = (profile.get("experience_details") or {}).get("roles") or []
    for ri, role in enumerate(roles, 2):
        data_cell(ws4, ri, 1, role.get("title") or "", bold=True)
        data_cell(ws4, ri, 2, role.get("company") or "")
        data_cell(ws4, ri, 3, role.get("start") or "")
        data_cell(ws4, ri, 4, role.get("end") or "")
        data_cell(ws4, ri, 5, role.get("months") or "")

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf
